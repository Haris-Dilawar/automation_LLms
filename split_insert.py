import regex as re
from docx import Document

def make_changes(string,word_path,row_index):
    # Split the text on "Assessment"
    parts = string.split('Assessment')
    # The first part is the "Introduction" section
    string1 = parts[0].strip()
    # The second part is the "Assessment" section
    string2 = parts[1].strip()
    # Remove the "Introduction" heading
    string1 = string1.replace('Introduction\n', '')
    string2 = string2.replace('The character of the new use and of the wider area\n', '')
    string2 = string2.replace('The size of the property\n', '')
    string2 = string2.replace('The nature and character of any services provided\n', '')
    string2 = string2.replace('The pattern of activity associated with the use including numbers of occupants, the period of use, issues of noise, disturbance, and parking demand\n', '')
    string2 = string2.replace('Complaints_Details: No complaints have been reported regarding the property.', '')
    
    # Split the text into lines for each section
    line_1 = string1.split('\n')
    line_2 = string2.split('\n')
    # Remove the numbering from each line
    line_1 = [re.sub(r'^\d+\.\d+\.\d+\.\s*', '', line) for line in line_1]
    line_2 = [re.sub(r'^\d+\.\d+\.\d+\.\s*', '', line) for line in line_2]
    line_2 = line_2[:-1]
    


    ### reading the word document
    doc = Document(word_path)
    changes = {
    30: line_1[0],
    32: line_1[1],
    34: line_1[2],
    36: line_1[3],
    171:line_2[0],
    173:line_2[1],
    175:line_2[2],
    179:line_2[3],
    181:line_2[4],
    183:line_2[5],
    185:line_2[6],
    189:line_2[7],
    191:line_2[8],
    195:line_2[9],
    197:line_2[10],
    199:line_2[11],
    201:line_2[12],
    205:line_2[13],
    207:line_2[14],
    209:line_2[15],

    # Add more as needed
    
    }

    

    # Loop over the dictionary
    for index, new_text in changes.items():
        # Clear all runs from the paragraph
        for run in doc.paragraphs[index].runs:
            run.text = ''
        
        # Add new text to the paragraph
        doc.paragraphs[index].add_run(new_text)
        

    # Save the final document with all changes
    doc.save(f'your_document_{row_index}_final.docx')


